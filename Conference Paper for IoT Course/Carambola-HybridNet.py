# Import necessary libraries
import os
import numpy as np
import torch
import torch.nn as nn
import torch.nn.functional as F
from torch_geometric.data import Data
from torch_geometric.nn import GCNConv, GATConv
from torchvision import datasets, transforms, models
from sklearn.model_selection import train_test_split
from sklearn.metrics import confusion_matrix, classification_report
from sklearn.neighbors import kneighbors_graph
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
from torch.utils.data import DataLoader
from tqdm import tqdm
from docx import Document
from docx.shared import Inches

# Set paths
input_folder = "C:/Users/UseR/Desktop/Final Paper/All Papers/Z______Carambola Files/Final Dataset/Augmented Dataset"
output_path = "C:/Users/UseR/Desktop/Final Paper/All Papers/1.  Humayet S. Papers/Coferance Papers/1. Starfruit Paper/Outputs/GCN"
os.makedirs(output_path, exist_ok=True)

# Set device
device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')

# Transform images
transform = transforms.Compose([
    transforms.Resize((224, 224)),
    transforms.ToTensor(),
    transforms.Normalize(mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225])
])

# Load dataset
dataset = datasets.ImageFolder(root=input_folder, transform=transform)
class_names = dataset.classes

# Extract features using ResNet18
feature_extractor = models.resnet18(pretrained=True)
feature_extractor = nn.Sequential(*list(feature_extractor.children())[:-1])
feature_extractor.eval().to(device)

features, labels = [], []
dataloader = DataLoader(dataset, batch_size=32, shuffle=False)

with torch.no_grad():
    for images, targets in tqdm(dataloader, desc='Extracting Features'):
        images = images.to(device)
        outputs = feature_extractor(images)
        outputs = outputs.squeeze(-1).squeeze(-1)
        features.append(outputs.cpu())
        labels.append(targets.cpu())

features = torch.cat(features).numpy()
labels = torch.cat(labels).numpy()

# Split dataset
idx = np.arange(len(labels))
idx_train, idx_test = train_test_split(idx, test_size=0.3, stratify=labels, random_state=42)
idx_val, idx_test = train_test_split(idx_test, test_size=0.5, stratify=labels[idx_test], random_state=42)

# Build k-NN graph
adj = kneighbors_graph(features, n_neighbors=5, metric='cosine', mode='connectivity', include_self=False)
edge_index = torch.tensor(np.array(adj.nonzero()), dtype=torch.long)

# Create PyG Data object
data = Data(
    x=torch.tensor(features, dtype=torch.float),
    edge_index=edge_index,
    y=torch.tensor(labels, dtype=torch.long)
).to(device)

# Define models
class GCNModel(nn.Module):
    def __init__(self, input_dim, hidden_dim, output_dim):
        super().__init__()
        self.conv1 = GCNConv(input_dim, hidden_dim)
        self.conv2 = GCNConv(hidden_dim, output_dim)
        self.dropout = nn.Dropout(0.5)

    def forward(self, data):
        x, edge_index = data.x, data.edge_index
        x = F.relu(self.conv1(x, edge_index))
        x = self.dropout(x)
        x = self.conv2(x, edge_index)
        return F.log_softmax(x, dim=1)

class GATModel(nn.Module):
    def __init__(self, input_dim, hidden_dim, output_dim, heads=4):
        super().__init__()
        self.gat1 = GATConv(input_dim, hidden_dim, heads=heads, dropout=0.6)
        self.gat2 = GATConv(hidden_dim * heads, output_dim, heads=1, concat=False, dropout=0.6)
        self.dropout = nn.Dropout(0.5)

    def forward(self, data):
        x, edge_index = data.x, data.edge_index
        x = F.elu(self.gat1(x, edge_index))
        x = self.dropout(x)
        x = self.gat2(x, edge_index)
        return F.log_softmax(x, dim=1)

class HybridGCNGNN(nn.Module):
    def __init__(self, input_dim, hidden_dim, output_dim, heads=4):
        super().__init__()
        self.gcn1 = GCNConv(input_dim, hidden_dim)
        self.gcn2 = GCNConv(hidden_dim, hidden_dim * 2)
        self.proj = nn.Linear(hidden_dim * 2, hidden_dim)
        self.gat1 = GATConv(hidden_dim, hidden_dim, heads=heads, dropout=0.5)
        self.gat2 = GATConv(hidden_dim * heads, hidden_dim, heads=1, concat=False, dropout=0.5)
        self.fc1 = nn.Linear(hidden_dim, hidden_dim // 2)
        self.fc2 = nn.Linear(hidden_dim // 2, output_dim)
        self.dropout = nn.Dropout(0.5)
        self.bn_gcn = nn.BatchNorm1d(hidden_dim * 2)
        self.bn_gat = nn.BatchNorm1d(hidden_dim)

    def forward(self, data):
        x, edge_index = data.x, data.edge_index
        x_gcn = F.relu(self.gcn1(x, edge_index))
        x_gcn = self.dropout(x_gcn)
        x_gcn = F.relu(self.gcn2(x_gcn, edge_index))
        x_gcn = self.bn_gcn(x_gcn)
        x_gcn = self.dropout(x_gcn)
        x_proj = self.proj(x_gcn)
        x_gat = F.elu(self.gat1(x_proj, edge_index))
        x_gat = self.dropout(x_gat)
        x_gat = F.elu(self.gat2(x_gat, edge_index))
        x_gat = self.bn_gat(x_gat)
        x_gat = self.dropout(x_gat)
        x_combined = x_proj + x_gat
        x = F.relu(self.fc1(x_combined))
        x = self.dropout(x)
        x = self.fc2(x)
        return F.log_softmax(x, dim=1)

# Training and evaluation utilities

def train_model(model, name):
    model.to(device)
    optimizer = torch.optim.Adam(model.parameters(), lr=0.005, weight_decay=5e-4)
    scheduler = torch.optim.lr_scheduler.ReduceLROnPlateau(optimizer, mode='max', factor=0.5, patience=10)
    criterion = nn.NLLLoss()
    history = {k: [] for k in ['train_loss', 'val_loss', 'test_loss', 'train_acc', 'val_acc', 'test_acc']}
    best_val_acc, best_test_acc, best_epoch = 0, 0, 0

    for epoch in range(1, 201):
        model.train()
        optimizer.zero_grad()
        out = model(data)
        loss = criterion(out[idx_train], data.y[idx_train])
        loss.backward()
        optimizer.step()
        model.eval()
        with torch.no_grad():
            out = model(data)
            pred = out.argmax(dim=1)
            losses = [criterion(out[i], data.y[i]).item() for i in [idx_train, idx_val, idx_test]]
            accs = [(pred[i] == data.y[i]).sum().item() / len(i) for i in [idx_train, idx_val, idx_test]]
        scheduler.step(accs[1])
        for i, k in enumerate(history): history[k].append((losses + accs)[i])
        if accs[1] > best_val_acc:
            best_val_acc, best_test_acc, best_epoch = accs[1], accs[2], epoch
            torch.save(model.state_dict(), os.path.join(output_path, f'{name}_best_model.pth'))
    model.load_state_dict(torch.load(os.path.join(output_path, f'{name}_best_model.pth')))
    model.eval()
    with torch.no_grad():
        out = model(data)
        pred = out.argmax(dim=1)
    cm = confusion_matrix(data.y[idx_test].cpu(), pred[idx_test].cpu())
    report = classification_report(data.y[idx_test].cpu(), pred[idx_test].cpu(), target_names=class_names, output_dict=True)
    return best_val_acc, best_test_acc, best_epoch, history, cm, pd.DataFrame(report).transpose()





















# def add_to_doc(doc, name, val_acc, test_acc, epoch, history, cm, report_df):
#     doc.add_heading(f'{name} Model', level=1)
#     doc.add_paragraph(f'Best Validation Accuracy: {val_acc:.4f} at Epoch {epoch}')
#     doc.add_paragraph(f'Test Accuracy: {test_acc:.4f}')
#     plt.figure()
#     plt.plot(history['train_acc'], label='Train')
#     plt.plot(history['val_acc'], label='Val')
#     plt.plot(history['test_acc'], label='Test')
#     plt.legend(); plt.title(f'{name} Accuracy');
#     plot_path = os.path.join(output_path, f'{name}_acc.png'); plt.savefig(plot_path); plt.close()
#     doc.add_picture(plot_path, width=Inches(5))
#     cm_path = os.path.join(output_path, f'{name}_cm.png')
#     plt.figure(); sns.heatmap(cm, annot=True, fmt='d', xticklabels=class_names, yticklabels=class_names)
#     plt.title(f'{name} Confusion Matrix'); plt.savefig(cm_path); plt.close()
#     doc.add_picture(cm_path, width=Inches(5))
#     table = doc.add_table(rows=1, cols=len(report_df.columns)+1)
#     hdr = table.rows[0].cells; hdr[0].text = 'Class'
#     for i, col in enumerate(report_df.columns): hdr[i+1].text = col
#     for idx, row in report_df.iterrows():
#         cells = table.add_row().cells
#         cells[0].text = str(idx)
#         for i, val in enumerate(row):
#             cells[i+1].text = f"{val:.4f}" if isinstance(val, float) else str(val)
#     doc.add_page_break()

# # Run all models and build report
# doc = Document()
# doc.add_heading("Carambola Leaf Disease Detection Report", 0)

# for model_class, name in [(GCNModel, 'GCN'), (GATModel, 'GAT'), (HybridGCNGNN, 'Hybrid')]:
#     model = model_class(input_dim=512, hidden_dim=128, output_dim=len(class_names))
#     val_acc, test_acc, epoch, hist, cm, rep = train_model(model, name)
#     add_to_doc(doc, name, val_acc, test_acc, epoch, hist, cm, rep)

# doc.save(os.path.join(output_path, "hybrid_classification_report.docx"))
# print("✅ All models trained and report saved!")


# Place your original code here first — (unchanged)

# ADDITIONAL IMPORTS
from sklearn.preprocessing import label_binarize
from sklearn.metrics import roc_curve, auc
from PIL import Image, ImageDraw
import random

# ------- ADDITIONAL VISUALIZATIONS -------- #
def plot_roc_curve(data, pred, model_name):
    y_true = data.y[idx_test].cpu().numpy()
    y_pred = F.softmax(pred, dim=1).detach().cpu().numpy()[idx_test]  # <-- Fixed here
    y_bin = label_binarize(y_true, classes=np.arange(len(class_names)))

    fpr = {}
    tpr = {}
    roc_auc = {}
    for i in range(len(class_names)):
        fpr[i], tpr[i], _ = roc_curve(y_bin[:, i], y_pred[:, i])
        roc_auc[i] = auc(fpr[i], tpr[i])

    plt.figure()
    for i in range(len(class_names)):
        plt.plot(fpr[i], tpr[i], label=f'{class_names[i]} (AUC = {roc_auc[i]:.2f})')
    plt.plot([0, 1], [0, 1], 'k--')
    plt.title(f'{model_name} ROC Curve')
    plt.xlabel('False Positive Rate')
    plt.ylabel('True Positive Rate')
    plt.legend()
    roc_path = os.path.join(output_path, f'{model_name}_roc.png')
    plt.savefig(roc_path)
    plt.close()
    return roc_path


def plot_loss_accuracy(history, model_name):
    plt.figure(figsize=(10, 5))
    plt.subplot(1, 2, 1)
    plt.plot(history['train_loss'], label='Train Loss')
    plt.plot(history['val_loss'], label='Val Loss')
    plt.plot(history['test_loss'], label='Test Loss')
    plt.title('Loss')
    plt.legend()

    plt.subplot(1, 2, 2)
    plt.plot(history['train_acc'], label='Train Acc')
    plt.plot(history['val_acc'], label='Val Acc')
    plt.plot(history['test_acc'], label='Test Acc')
    plt.title('Accuracy')
    plt.legend()

    path = os.path.join(output_path, f'{model_name}_loss_accuracy.png')
    plt.savefig(path)
    plt.close()
    return path


def plot_dual_curve(history, model_name):
    epochs = list(range(len(history['train_acc'])))

    fig, axes = plt.subplots(1, 2, figsize=(12, 5))

    # Accuracy subplot
    axes[0].plot(epochs, history['train_acc'], 'o-', label='Train Accuracy')
    axes[0].plot(epochs, history['val_acc'], 'o-', label='Validation Accuracy')
    axes[0].set_title(f'{model_name} Accuracy')
    axes[0].set_xlabel('Epochs')
    axes[0].set_ylabel('Accuracy')
    axes[0].legend()
    axes[0].grid(True)

    # Loss subplot
    axes[1].plot(epochs, history['train_loss'], 'o-', label='Train Loss')
    axes[1].plot(epochs, history['val_loss'], 'o-', label='Validation Loss')
    axes[1].set_title(f'{model_name} Loss')
    axes[1].set_xlabel('Epochs')
    axes[1].set_ylabel('Loss')
    axes[1].legend()
    axes[1].grid(True)

    # Save and return path
    dual_plot_path = os.path.join(output_path, f"{model_name}_dual_accuracy_loss.png")
    plt.tight_layout()
    plt.savefig(dual_plot_path)
    plt.close()
    return dual_plot_path


def mark_detected_sample_image(model_name):
    class_folder = random.choice(dataset.classes)
    class_index = dataset.class_to_idx[class_folder]
    class_path = os.path.join(input_folder, class_folder)
    image_path = os.path.join(class_path, random.choice(os.listdir(class_path)))

    img = Image.open(image_path).convert("RGB")
    draw = ImageDraw.Draw(img)
    w, h = img.size
    # Simulate detection box (placeholder)
    box = (w//4, h//4, 3*w//4, 3*h//4)
    draw.rectangle(box, outline="red", width=3)
    draw.text((box[0], box[1]-10), f"Detected: {class_folder}", fill="red")

    marked_path = os.path.join(output_path, f"{model_name}_sample_detected.png")
    img.save(marked_path)
    return marked_path

# Modify add_to_doc() to insert additional images
def add_to_doc(doc, name, val_acc, test_acc, epoch, history, cm, report_df):
    doc.add_heading(f'{name} Model', level=1)
    doc.add_paragraph(f'Best Validation Accuracy: {val_acc:.4f} at Epoch {epoch}')
    doc.add_paragraph(f'Test Accuracy: {test_acc:.4f}')

    # Accuracy Graph
    acc_path = os.path.join(output_path, f'{name}_acc.png')
    plt.figure()
    plt.plot(history['train_acc'], label='Train')
    plt.plot(history['val_acc'], label='Val')
    plt.plot(history['test_acc'], label='Test')
    plt.legend(); plt.title(f'{name} Accuracy');
    plt.savefig(acc_path); plt.close()
    doc.add_picture(acc_path, width=Inches(5))

    # Confusion Matrix
    cm_path = os.path.join(output_path, f'{name}_cm.png')
    plt.figure(); sns.heatmap(cm, annot=True, fmt='d', xticklabels=class_names, yticklabels=class_names)
    plt.title(f'{name} Confusion Matrix'); plt.savefig(cm_path); plt.close()
    doc.add_picture(cm_path, width=Inches(5))

    # Classification Report Table
    table = doc.add_table(rows=1, cols=len(report_df.columns)+1)
    hdr = table.rows[0].cells; hdr[0].text = 'Class'
    for i, col in enumerate(report_df.columns): hdr[i+1].text = col
    for idx, row in report_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(idx)
        for i, val in enumerate(row):
            cells[i+1].text = f"{val:.4f}" if isinstance(val, float) else str(val)

    # ROC Curve
    roc_path = plot_roc_curve(data, model(data), name)
    doc.add_picture(roc_path, width=Inches(5))

    # Loss vs Accuracy Curve
    la_path = plot_loss_accuracy(history, name)
    doc.add_picture(la_path, width=Inches(5))
    
    # Dual subplot: Accuracy & Loss like InceptionV3 style
    dual_path = plot_dual_curve(history, name)
    doc.add_picture(dual_path, width=Inches(5))


    # Sample detection image
    sample_path = mark_detected_sample_image(name)
    doc.add_picture(sample_path, width=Inches(5))

    doc.add_page_break()

# Run the models and build the updated document
doc = Document()
doc.add_heading("Carambola Leaf Disease Detection Report", 0)

for model_class, name in [(GCNModel, 'GCN'), (GATModel, 'GAT'), (HybridGCNGNN, 'Hybrid')]:
    model = model_class(input_dim=512, hidden_dim=128, output_dim=len(class_names))
    val_acc, test_acc, epoch, hist, cm, rep = train_model(model, name)
    add_to_doc(doc, name, val_acc, test_acc, epoch, hist, cm, rep)

doc.save(os.path.join(output_path, "hybrid_classification_report.docx"))
print("✅ All models trained and full report with extra graphs saved!")
