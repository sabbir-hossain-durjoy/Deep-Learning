import os
import numpy as np
import tensorflow as tf
import matplotlib.pyplot as plt
from tensorflow.keras.preprocessing.image import ImageDataGenerator
from tensorflow.keras.applications import InceptionV3
from tensorflow.keras.layers import Dense, GlobalAveragePooling2D, Dropout
from tensorflow.keras.models import Model
from tensorflow.keras.optimizers import Adam
from tensorflow.keras.callbacks import ModelCheckpoint
from sklearn.metrics import classification_report, confusion_matrix
import seaborn as sns
from docx import Document
from docx.shared import Inches

input_path = r""
output_path = r""
os.makedirs(output_path, exist_ok=True)

EPOCHS = 30
BATCH_SIZE = 32
LEARNING_RATE = 0.0001

doc = Document()
doc.add_heading('Classification Report for InceptionV3', 0)

def create_model(input_shape, num_classes):
    base_model = InceptionV3(weights='imagenet', include_top=False, input_shape=input_shape)
    base_model.trainable = False
    x = base_model.output
    x = GlobalAveragePooling2D()(x)
    x = Dense(1024, activation='relu')(x)
    x = Dropout(0.5)(x)
    predictions = Dense(num_classes, activation='softmax')(x)
    model = Model(inputs=base_model.input, outputs=predictions)
    return model

def process_data(input_path, img_size, batch_size):
    train_datagen = ImageDataGenerator(
        preprocessing_function=tf.keras.applications.inception_v3.preprocess_input,
        validation_split=0.2,
        rotation_range=30,
        width_shift_range=0.2,
        height_shift_range=0.2,
        shear_range=0.2,
        zoom_range=0.2,
        horizontal_flip=True,
        fill_mode="nearest"
    )
    train_generator = train_datagen.flow_from_directory(
        input_path,
        target_size=img_size,
        batch_size=batch_size,
        class_mode='categorical',
        subset='training',
        shuffle=True
    )
    validation_generator = train_datagen.flow_from_directory(
        input_path,
        target_size=img_size,
        batch_size=batch_size,
        class_mode='categorical',
        subset='validation',
        shuffle=False
    )
    return train_generator, validation_generator

model_name = 'InceptionV3'
img_size = (299, 299)

train_generator, validation_generator = process_data(input_path, img_size, BATCH_SIZE)

model = create_model((img_size[0], img_size[1], 3), train_generator.num_classes)

model.compile(optimizer=Adam(learning_rate=LEARNING_RATE), 
              loss='categorical_crossentropy', 
              metrics=['accuracy'])

checkpoint = ModelCheckpoint(os.path.join(output_path, f'{model_name}_best_model.keras'),
                             monitor='val_accuracy', 
                             save_best_only=True, 
                             mode='max', 
                             verbose=1)

history = model.fit(
    train_generator,
    epochs=EPOCHS,
    validation_data=validation_generator,
    callbacks=[checkpoint]
)

plt.figure(figsize=(12, 5))
plt.subplot(1, 2, 1)
plt.plot(history.history['accuracy'], label='Train Accuracy', marker='o')
plt.plot(history.history['val_accuracy'], label='Validation Accuracy', marker='o')
plt.title(f"{model_name} Accuracy")
plt.xlabel("Epochs")
plt.ylabel("Accuracy")
plt.legend()

plt.subplot(1, 2, 2)
plt.plot(history.history['loss'], label='Train Loss', marker='o')
plt.plot(history.history['val_loss'], label='Validation Loss', marker='o')
plt.title(f"{model_name} Loss")
plt.xlabel("Epochs")
plt.ylabel("Loss")
plt.legend()

accuracy_loss_curve_path = os.path.join(output_path, f'{model_name}_accuracy_loss_curve.png')
plt.savefig(accuracy_loss_curve_path)
plt.close()

model.load_weights(os.path.join(output_path, f'{model_name}_best_model.keras'))

eval_result = model.evaluate(validation_generator)

Y_pred = model.predict(validation_generator)
y_pred = np.argmax(Y_pred, axis=1)
y_true = validation_generator.classes

report = classification_report(y_true, y_pred, target_names=list(train_generator.class_indices.keys()))

cm = confusion_matrix(y_true, y_pred)
confusion_matrix_path = os.path.join(output_path, f'{model_name}_confusion_matrix.png')

plt.figure(figsize=(12, 10))
sns.heatmap(cm, annot=True, fmt='d', cmap='Blues',
            xticklabels=train_generator.class_indices.keys(),
            yticklabels=train_generator.class_indices.keys(),
            annot_kws={"size": 14})
plt.xticks(rotation=45, fontsize=12)
plt.yticks(rotation=0, fontsize=12)
plt.xlabel('Predicted', fontsize=14)
plt.ylabel('Actual', fontsize=14)
plt.title(f'{model_name} Confusion Matrix', fontsize=16)
plt.tight_layout()
plt.savefig(confusion_matrix_path)
plt.close()

doc.add_heading(f'{model_name} Performance', level=1)
doc.add_paragraph(f"Final Validation Loss: {eval_result[0]:.4f}")
doc.add_paragraph(f"Final Validation Accuracy: {eval_result[1]:.4f}")

doc.add_heading(f'{model_name} Classification Metrics', level=2)
doc.add_paragraph(report)

doc.add_heading(f'{model_name} Training History', level=2)
doc.add_picture(accuracy_loss_curve_path, width=Inches(6))

doc.add_heading(f'{model_name} Confusion Matrix', level=2)
doc.add_picture(confusion_matrix_path, width=Inches(6))

doc.save(os.path.join(output_path, f'{model_name}_classification_report.docx'))

print(f"{model_name} training completed and results saved.")
