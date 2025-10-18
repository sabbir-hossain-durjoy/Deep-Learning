import os
import numpy as np
import tensorflow as tf
import matplotlib.pyplot as plt
from tensorflow.keras.preprocessing.image import ImageDataGenerator
from tensorflow.keras.applications import InceptionV3, EfficientNetB0, DenseNet121, MobileNet, ResNet50, Xception, VGG16, VGG19
from tensorflow.keras.layers import Dense, GlobalAveragePooling2D, Dropout
from tensorflow.keras.models import Model
from tensorflow.keras.optimizers import Adam
from tensorflow.keras.callbacks import ModelCheckpoint
from sklearn.metrics import classification_report, confusion_matrix
import seaborn as sns
from docx import Document
from docx.shared import Inches

# Paths for input data and output result
input_path = r""
output_path = r""
os.makedirs(output_path, exist_ok=True)

# Parameters
EPOCHS = 30
BATCH_SIZE = 32
LEARNING_RATE = 0.0001  # Adjusted learning rate

# Create Document for Results
doc = Document()
doc.add_heading('Classification Report for All Models', 0)

# Function to create a model based on selected architecture
def create_model(model_name, input_shape, num_classes):
    base_model = None
    if model_name == 'InceptionV3':
        base_model = InceptionV3(weights='imagenet', include_top=False, input_shape=input_shape)
    elif model_name == 'EfficientNetB0':
        base_model = EfficientNetB0(weights='imagenet', include_top=False, input_shape=input_shape)
    elif model_name == 'DenseNet121':
        base_model = DenseNet121(weights='imagenet', include_top=False, input_shape=input_shape)
    elif model_name == 'MobileNet':
        base_model = MobileNet(weights='imagenet', include_top=False, input_shape=input_shape)
    elif model_name == 'ResNet50':
        base_model = ResNet50(weights='imagenet', include_top=False, input_shape=input_shape)
    elif model_name == 'Xception':
        base_model = Xception(weights='imagenet', include_top=False, input_shape=input_shape)
    elif model_name == 'VGG16':
        base_model = VGG16(weights='imagenet', include_top=False, input_shape=input_shape)
    elif model_name == 'VGG19':
        base_model = VGG19(weights='imagenet', include_top=False, input_shape=input_shape)
    
    base_model.trainable = False  # Freeze base model

    # Custom Model Architecture
    x = base_model.output
    x = GlobalAveragePooling2D()(x)
    x = Dense(1024, activation='relu')(x)
    x = Dropout(0.5)(x)  # Dropout to prevent overfitting
    predictions = Dense(num_classes, activation='softmax')(x)

    model = Model(inputs=base_model.input, outputs=predictions)
    return model

# Function to process the data
def process_data(input_path, img_size, batch_size):
    train_datagen = ImageDataGenerator(
        preprocessing_function=tf.keras.applications.inception_v3.preprocess_input,  # Adjust preprocessing function
        validation_split=0.2,  
        rotation_range=30,
        width_shift_range=0.2,
        height_shift_range=0.2,
        shear_range=0.2,
        zoom_range=0.2,
        horizontal_flip=True,
        fill_mode="nearest"
    )
    
    # Train and Validation Generators
    train_generator = train_datagen.flow_from_directory(
        input_path,
        target_size=img_size,
        batch_size=batch_size,
        class_mode='categorical',  # Proper label format
        subset='training',
        shuffle=True
    )

    validation_generator = train_datagen.flow_from_directory(
        input_path,
        target_size=img_size,
        batch_size=batch_size,
        class_mode='categorical',
        subset='validation',
        shuffle=False
    )
    
    return train_generator, validation_generator

# List of Models and their Input Sizes
models_info = {
    'InceptionV3': (299, 299),
    # 'EfficientNetB0': (224, 224),
    'DenseNet121': (224, 224),
    # 'MobileNet': (224, 224),
    # 'ResNet50': (224, 224),
    # 'Xception': (299, 299),
    # 'VGG16': (224, 224),
    # 'VGG19': (224, 224)
}

# Train models one by one
for model_name, input_size in models_info.items():
    img_size = input_size
    print(f"Training {model_name}...")

    # Process Data for each model
    train_generator, validation_generator = process_data(input_path, img_size, BATCH_SIZE)

    # Create the Model
    model = create_model(model_name, (img_size[0], img_size[1], 3), train_generator.num_classes)

    # Compile the Model
    model.compile(optimizer=Adam(learning_rate=LEARNING_RATE), 
                  loss='categorical_crossentropy', 
                  metrics=['accuracy'])

    # Model Checkpoint
    checkpoint = ModelCheckpoint(os.path.join(output_path, f'{model_name}_best_model.keras'),
                                 monitor='val_accuracy', 
                                 save_best_only=True, 
                                 mode='max', 
                                 verbose=1)

    # Train the Model (No EarlyStopping to allow full 30 epochs)
    history = model.fit(
        train_generator,
        epochs=EPOCHS,
        validation_data=validation_generator,
        callbacks=[checkpoint]
    )

    # Plot Accuracy & Loss Curves
    plt.figure(figsize=(12, 5))

    # Plot Accuracy
    plt.subplot(1, 2, 1)
    plt.plot(history.history['accuracy'], label='Train Accuracy', marker='o')
    plt.plot(history.history['val_accuracy'], label='Validation Accuracy', marker='o')
    plt.title(f"{model_name} Accuracy")
    plt.xlabel("Epochs")
    plt.ylabel("Accuracy")
    plt.legend()

    # Plot Loss
    plt.subplot(1, 2, 2)
    plt.plot(history.history['loss'], label='Train Loss', marker='o')
    plt.plot(history.history['val_loss'], label='Validation Loss', marker='o')
    plt.title(f"{model_name} Loss")
    plt.xlabel("Epochs")
    plt.ylabel("Loss")
    plt.legend()

    accuracy_loss_curve_path = os.path.join(output_path, f'{model_name}_accuracy_loss_curve.png')
    plt.savefig(accuracy_loss_curve_path)
    plt.close()

    # Load Best Model for Evaluation
    model.load_weights(os.path.join(output_path, f'{model_name}_best_model.keras'))

    # Evaluate the Model
    eval_result = model.evaluate(validation_generator)

    # Generate Predictions and Classification Report
    Y_pred = model.predict(validation_generator)
    y_pred = np.argmax(Y_pred, axis=1)
    y_true = validation_generator.classes

    report = classification_report(y_true, y_pred, target_names=list(train_generator.class_indices.keys()))

    # Confusion Matrix
    cm = confusion_matrix(y_true, y_pred)
    confusion_matrix_path = os.path.join(output_path, f'{model_name}_confusion_matrix.png')

    plt.figure(figsize=(12, 10))
    sns.heatmap(cm, annot=True, fmt='d', cmap='Blues',
                xticklabels=train_generator.class_indices.keys(),
                yticklabels=train_generator.class_indices.keys(),
                annot_kws={"size": 14})  # Increased font size for annotation
    plt.xticks(rotation=45, fontsize=12)
    plt.yticks(rotation=0, fontsize=12)
    plt.xlabel('Predicted', fontsize=14)
    plt.ylabel('Actual', fontsize=14)
    plt.title(f'{model_name} Confusion Matrix', fontsize=16)
    plt.tight_layout()  # Ensures labels fit inside the plot
    plt.savefig(confusion_matrix_path)
    plt.close()

    # Add Results to the Document
    doc.add_heading(f'{model_name} Performance', level=1)
    doc.add_paragraph(f"Final Validation Loss: {eval_result[0]:.4f}")
    doc.add_paragraph(f"Final Validation Accuracy: {eval_result[1]:.4f}")

    # Classification Metrics Section
    doc.add_heading(f'{model_name} Classification Metrics', level=2)
    doc.add_paragraph(report)

    # Training History Section
    doc.add_heading(f'{model_name} Training History', level=2)
    doc.add_picture(accuracy_loss_curve_path, width=Inches(6))

    # Confusion Matrix Section
    doc.add_heading(f'{model_name} Confusion Matrix', level=2)
    doc.add_picture(confusion_matrix_path, width=Inches(6))

    print(f"{model_name} Training completed and results saved.")

# Save Document after all models are processed
doc.save(os.path.join(output_path, 'all_models_classification_report.docx'))
print("All models' training completed and results saved in Word document!")
